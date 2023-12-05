# Bring in deps
import os
import docx
import streamlit as st
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
import io
import requests

os.environ['OPENAI_API_KEY'] = "<API_KEY>"

# App framework
st.title('💡🔗 Concept Sprint GPT')
st.text('PoV by Srikar Kashyap Pulipaka')
prompt = st.text_input('Plug in your prompt here. Example: Mobile Homes')

# Prompt templates

prompt1_template = PromptTemplate(
    input_variables=['topic'],
    template="""
Background:
I am working on a project that aims to revolutionize {topic}, specifically targeting a FUTURISITC SYNDICATE. I am organizing a concept sprint workshop where stakeholders from various fields come together to brainstorm and strategize on the futuristic solutions for this sector.
Objective:
I want to create a synthetic simulation of the concept sprint workshop to explore potential solutions and approaches.
Task:
Please help me identify and create personalities for SIX stakeholder roles that would be instrumental in this project. These roles should cover a broad spectrum of expertise and perspectives, including:- Project VisionandBusiness Goals;Voice of the user and User research;Existing Product Audit and DesignEvolution;Competitor Audit;Technology Considerations and Opportunities;Comparable Solution in aDifferentSpace.
For each role, please provide the following information in the fixed format in markdown

Question:
Can you assist me in developing this stakeholder simulation and creating a robust conversation that culminates in a well-defined framework for my project? Additionally, could you guide me on how this simulation can betranslated into actionable tasks for a real-world use case?

Provide output in this fixed format in markdown:
-----------------------------------

**Name**: name

**Email**: email

**Business Area**: business area

**SME in**: SME in
"""
)

prompt2_template = PromptTemplate(
    input_variables=['personas', 'topic'],
    template="""
   Objective:
I want to create a synthetic simulation of the concept sprint workshop to explore potential solutions and approaches.
Background:
I am working on a project that aims to revolutionize {topic} sector specifically targeting a futuristic syndicate. I am organizing a concept sprint workshop where stakeholders from various fields come together to brainstorm and strategize on the futuristic solutions for this sector.
(Here is a summary of the outcomes of previous stages of the sprint)
I have identified and create personalities for  stakeholder roles that would be instrumental in this project. These roles cover a broad spectrum of expertise and perspectives and expertise.
{personas}


I would like to simulate a conversation between them where they discuss and come to a consensus on the following aspects of the project...
Sprint Challenge:
The group agrees on a broad statement knowing this may change as we learn more; list out the deliverables the group will create
[GOAL] +[USER] +[PLATFORM] +[TIMEFRAME]
Sprint Deliverables:

Goal and Anti-Goals
The group brainstorms and agrees on the top 3 goals and top 3 anti-goals.
Goals:
Describe the high-level objectives (not features).
Discuss and confirm if they apply to this sprint or a future sprint
What user or users will you focus on?
What key moments or pain points do you want to sketch around to have the most impact?
What does success look like?
How will you measure it?
Do you need any measurement tools?
What are the top 3 goals?
Non-Goals:
Document what are explicitly not a goal and may be an anti-goal
What are the top 3 anti-goals?
Anti-goals are explicit things that are not a goal and serve to direct the teams thinking
Assumptions Made (I):
With you long term goal in mind document the assumptions that were made
Assumptions:
List the underlying assumptions for the concept and prioritize the questions and data you still need to find out about
List all the assumptions that are underlying your concept
What primary data questions would you need to clarify these assumptions
Prioritize the most important questions you need answered.
Return properly formatted markdown in the following format:

-----------------------------------

### Conversation

<Conversation with atleast 5 exchanges>

### Goals

1. Goal 1

2. Goal 2

3. Goal 3

### Anti-Goals

1. Anti-Goal 1

2. Anti-Goal 2

3. Anti-Goal 3

### Assumptions

1. Assumption 1

2. Assumption 2

3. Assumption 3

### Questions

1. Question 1

2. Question 2

3. Question 3
"""
)

prompt3_template = PromptTemplate(
    input_variables=['personas', 'topic', 'conversation'],
    template="""
Objective:
I want to create a synthetic simulation of the concept sprint workshop to explore potential solutions and approaches.
Background:
I am working on a project that aims to revolutionize {topic} sector specifically targeting a futuristic syndicate. I am organizing a concept sprint workshop where stakeholders from various fields come together to brainstorm and strategize on the futuristic solutions for this sector.
(Here is a summary of the outcomes of previous stages of the sprint)
I have identified and create personalities for  stakeholder roles that would be instrumental in this project. These roles cover a broad spectrum of expertise and perspectives and expertise.
{personas}
After a conversation, the following goals, anti-goals and assumptions have been arrived at:
{conversation}

I would like to simulate a conversation between them where they discuss and come to a consensus on the following aspects of the project...

Allocate each member of the group a lightening talk topic (see topics below) to present. Each talk should be around 4-5 sentences and of TED talk format.
each. Topics include vision/goals, voice of user, existing produces, competitor v iew, technology available.
-Project Vision / Business Goals
-Voice of the user / User research
-Ex isting Product Audit / Design Ev olution
-Competitor Audit
-Technology Considerations & Opportunities
-Comparable Solution in a Different Space

At the end of the talks, the group generates 15 How Might We (HMW) based on the information provided in the talks.

Now we will group the HMW post-it's into common categories, similar ideas, and overlapping concepts. 
Then each person has 3 votes for selecting the best 3 ideas. Look for the ideas or clusters with the most 
votes. After voting, present output in the following fixed markdown format:

-----------------------------------

### Lightening Talks

1. Talk 1
2. Talk 2
3. Talk 3
4. Talk 4
5. Talk 5
6. Talk 6

### List of HMW

1. 15 bulleted HMWs with their category

### HMWs sorted by number of votes received descending

1. HMW 1
2. HMW 2
3. HMW 3

"""
)

prompt4_template = PromptTemplate(
    input_variables=['topic', 'conversation', 'hmws', 'personas'],
    template="""
Objective:
I want to create a synthetic simulation of the concept sprint workshop to explore potential solutions and approaches.
Background:
I am working on a project that aims to revolutionize {topic} sector specifically targeting a futuristic syndicate. I am organizing a concept sprint workshop where stakeholders from various fields come together to brainstorm and strategize on the futuristic solutions for this sector.
(Here is a summary of the outcomes of previous stages of the sprint)
I have identified and create personalities for  stakeholder roles that would be instrumental in this project. These roles cover a broad spectrum of expertise and perspectives and expertise.
{personas}
After a conversation, the following goals, anti-goals and assumptions have been arrived at:
{conversation}
The following HMWs have been finalized after lighting talks and voting:
{hmws}
I would like to simulate a conversation between them where they discuss and come to a consensus on the following aspects of the project...
Crazy 8's and Vote:
Each sprinter will now sketch out 8 ideas for solving/attacking the concept. Work alone for 8 minutes and 
don't overthink. Then we share and vote on the solutions. It can be beneficial to run multiple rounds of Crazy 
8's.
Select and/or Combine:
If there are more than one w inning concepts look for unique ways to combine the into one. Once decided 
sketch out the final concept into a self-standing that any one could pick up and digest
Assumptions Made (II):
With you winning concept in mind document the assumptions that were made

Select the top 10 ideas after voting and present output in the following fixed markdown format:
-----------------------------------
### Crazy 8's

1. Ten Ideas in this format

"""
)


prompt5_template = PromptTemplate(
    input_variables=['topic', 'conversation', 'hmws', 'personas'],
    template="""
Given the below solution and ideas
{hmws}
Build me a Wardley map for the above solution, and outcomes provided. Utilize onlinewardleymapping.com which uses a mapping language described below;
To set the title
Example:
title My Wardley Map
------------------------
To create a component
component Name [Visibility (Y Axis), Maturity (X Axis)]
Example:
component Customer [0.9, 0.5]
component Cup of Tea [0.9, 0.5]
------------------------
Inertia - component likely to face resistance to change.
component Name [Visibility (Y Axis), Maturity (X Axis)] inertia
Example:
component Customer [0.95, 0.5] inertia
component Cup of Tea [0.9, 0
------------------------
To evolve a component
evole Name (X Axis)
Example:
evolve Customer 0.8
evolve Cup of Tea evolve 0.8
------------------------
To link components
Example:
Start Component->End Component
Customer->Cup of Tea
------------------------
To set component as pipeline:
pipeline Component Name [X Axis (start), X Axis (end)]
Example:
pipeline Customer [0.15, 0.9]
pipeline Customer
------------------------
"""
)
# define the LLM and the various prompts
llm = OpenAI(temperature=0, max_tokens=4800, model_name='gpt-4')

prompt1_chain = LLMChain(llm=llm, prompt=prompt1_template,
                         verbose=True, output_key='personas',)

prompt2_chain = LLMChain(llm=llm, prompt=prompt2_template,
                         verbose=True, output_key='conversation')

prompt3_chain = LLMChain(llm=llm, prompt=prompt3_template,
                            verbose=True, output_key='hmws')

prompt4_chain = LLMChain(llm=llm, prompt=prompt4_template,
                            verbose=True, output_key='report')

prompt5_chain = LLMChain(llm=llm, prompt=prompt5_template,
                            verbose=True, output_key='wardley_map')
if "file_created" not in st.session_state:
    st.session_state.file_created = False

def download_report(prompt, personas, conversation, hmws, report):
    """
    Download the report as a docx file
    :param prompt: prompt
    :param personas: personas
    :param conversation: conversation
    :param hmws: hmws
    :param report: report
    :return: docx file
    """

    print('Downloading report')
    doc = docx.Document()
    doc.add_heading('Concept Sprint Report', 0)
    doc.add_heading('Author: Srikar Kashyap Pulipaka', 3)
    doc.add_heading('Prompt', 1)
    doc.add_paragraph(prompt)
    doc.add_heading('Personas', 1)
    doc.add_paragraph(personas)
    doc.add_heading('Conversation', 1)
    doc.add_paragraph(conversation)
    doc.add_heading('HMWs', 1)
    doc.add_paragraph(hmws)
    doc.add_heading('Crazy 8s', 1)
    doc.add_paragraph(report)

    # return doc link to download
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def save_wardley_map(url, data):
    """
    Save a wardley map to the online wardley map tool
    :param url: url to the online wardley map tool
    :param data: data to be sent to the online wardley map tool
    :return: response from the online wardley map tool
    """
    headers = {'Content-Type': 'application/json'}
    response = requests.post(url, headers=headers, json=data)

    if response.status_code == 200:
        return response.json()
    else:
        return response.status_code


if prompt:
    st.markdown('Here is your prompt:')
    st.markdown(prompt)
    with st.spinner('Generating personas...'):
        personas = prompt1_chain.run(topic=prompt,)
    st.markdown('-----------------------------------')
    st.markdown('Here are the personas:')
    st.markdown(personas)
    st.markdown('-----------------------------------')
    with st.spinner('Generating the conversation....'):
        conversation = prompt2_chain.run(personas=personas, topic=prompt)
    # st.markdown('-----------------------------------')
    st.markdown('Here is the conversation:')
    st.markdown(conversation)
    st.markdown('-----------------------------------')
    with st.spinner('Generating the Lightening talk and HMWs....'):
        hmws = prompt3_chain.run(personas=personas, topic=prompt, conversation=conversation)
    st.markdown('Here is the lightening talks and HMWs:')
    st.markdown(hmws)
    st.markdown('-----------------------------------')
    top3_hmws = hmws.split('### HMWs sorted by number of votes received descending')[1]
    with st.spinner('Generating Crazy 8s and ideas....'):
        report = prompt4_chain.run(topic=prompt, conversation=conversation, hmws=top3_hmws, personas=personas)
    st.markdown('Here is the Final concept and assumptions made:')
    st.markdown(report)
    st.markdown('-----------------------------------')
    with st.spinner('Generating the wardley map....'):
        wardley_map = prompt5_chain.run(topic=prompt, conversation=conversation, hmws=report, personas=personas)
    st.markdown('Here is the wardley map:')
    st.markdown(wardley_map)
    file_data = download_report(prompt, personas, conversation, hmws, report)
    wardley_link = save_wardley_map("https://api.onlinewardleymaps.com/v1/maps/save", {"text": wardley_map})
    st.markdown('-----------------------------------')
    st.markdown('Here is the link to the wardley map:')
    wardley_id = wardley_link['id']
    wardley_link = f"https://onlinewardleymaps.com/#{wardley_id}"
    st.markdown(wardley_link)
    st.download_button(label='Download Report', data=file_data, file_name='Concept_Sprint_Report.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',)
   

